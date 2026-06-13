from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/AI_Native新人成长教练_方案说明.docx")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix) :])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, size=10, bold=bold)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def main():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("AI Native 新人成长教练方案说明")
    set_run_font(r, size=18, bold=True, color="1F4D78")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("面向腾讯游戏部门新人及转 AI 员工的 90 天学习与融入 MVP")
    set_run_font(r, size=10, color="667085")

    doc.add_heading("一、问题诊断", level=1)
    add_para(
        doc,
        "AI Native 组织中，新人不仅要熟悉制度和业务，还要快速建立 AI 协作思维。传统入职培训偏宣讲，缺少岗位化场景、能力量化、持续练习和成果沉淀，导致学习难转化为真实业务价值。",
    )

    doc.add_heading("二、方案设计", level=1)
    add_para(
        doc,
        "本方案设计为一个可运行网页 Demo：AI Native 新人成长教练。用户选择策划、美术、运营、HR、产品岗位后，通过 AI 教练完成 10 个轻量诊断问题，生成 AI 等级、五维能力雷达图、短板分析和 30/60/90 天成长路线。",
    )
    add_bullet(doc, "AI 教练：承接诊断、GAP 分析、场景发现和成长路线生成。")
    add_bullet(doc, "学习广场：按岗位推荐微课、知识卡片和业务练习。")
    add_bullet(doc, "实践社区：沉淀最佳实践、Prompt、工作流和 AI 案例。")
    add_bullet(doc, "激励机制：完成诊断、学习、练习、发布案例可获得 XP，成长伙伴随 XP 升级。")

    doc.add_heading("三、AI 工具选型理由", level=1)
    add_para(
        doc,
        "MVP 阶段采用规则引擎与 Mock 数据，优先保证演示闭环稳定、无需账号和 API Key。系统已预留 AI Provider 接口，后续可替换为 OpenAI 或企业私有模型，用于实时诊断、学习计划生成和练习点评。",
    )

    doc.add_heading("四、关键配置", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.7)
    table.columns[1].width = Inches(4.8)
    headers = table.rows[0].cells
    set_cell_text(headers[0], "配置项", True)
    set_cell_text(headers[1], "说明", True)
    shade(headers[0], "F2F4F7")
    shade(headers[1], "F2F4F7")
    rows = [
        ("能力模型", "AI认知、使用习惯、Prompt能力、工作流能力、创新能力，每项 0-20 分。"),
        ("岗位范围", "策划、美术、运营、HR、产品，学习资源和实践任务按岗位生成。"),
        ("部署方式", "静态网页，可部署到 Vercel、Netlify、GitHub Pages 或企业内部门户。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left, True)
        set_cell_text(cells[1], right)

    doc.add_heading("五、迭代记录与效果评估", level=1)
    add_para(
        doc,
        "第一版先完成可运行闭环：岗位选择、诊断、计划、学习、练习点评、社区发布和 XP 反馈。后续迭代将接入真实模型、员工数据看板和团队案例审核。评估指标包括诊断完成率、学习计划生成率、练习完成率、社区分享率，以及完成真实业务 AI 实践的用户占比。",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
